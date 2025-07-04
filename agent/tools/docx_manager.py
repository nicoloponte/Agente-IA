# agent/tools/docx_manager.py
from langchain.tools import BaseTool
from docx import Document

class DocxManagerTool(BaseTool):
    name: str = "Gestionar documentos Docx"
    description: str = "Útil para crear y editar documentos Docx."

    def _run(self, command: str) -> str:
        try:
            parts = command.split(" ", 2)
            action = parts[0]
            file_path = parts[1]

            if action == "crear":
                doc = Document()
                doc.save(file_path)
                return f"Documento Docx {file_path} creado."
            elif action == "escribir":
                content = parts[2]
                doc = Document(file_path)
                doc.add_paragraph(content)
                doc.save(file_path)
                return f"Contenido escrito en {file_path}."
            else:
                return "Acción no válida."
        except Exception as e:
            return str(e)

    async def _arun(self, command: str) -> str:
        raise NotImplementedError("No implementado")